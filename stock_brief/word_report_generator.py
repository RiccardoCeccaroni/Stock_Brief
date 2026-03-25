"""
word_report_generator.py
Generates a Word (.docx) investment research report using python-docx.
"""

import os
import re
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from report_generator import _fmt_number


def _normalize_markdown(text: str) -> str:
    """Convert markdown headings (### Foo) to bold (**Foo**)."""
    lines = text.split("\n")
    result = []
    for line in lines:
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            heading_text = m.group(2).strip()
            if not heading_text.startswith("**"):
                heading_text = f"**{heading_text}**"
            result.append(heading_text)
        else:
            result.append(line)
    return "\n".join(result)


def _add_section_heading(doc, text):
    """Add a styled section heading."""
    para = doc.add_heading(text, level=2)
    for run in para.runs:
        run.font.color.rgb = RGBColor(30, 30, 30)
        run.font.size = Pt(14)


def _set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)


def _style_header_row(row, headers):
    """Style a table header row with dark background and white text."""
    for i, cell in enumerate(row.cells):
        cell.text = headers[i]
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "283C78")


def generate_word_report(
    profile: dict,
    description: str,
    business_model_section: str,
    income_statements: list,
    cash_flows: list,
    key_metrics: list,
    peer_comparison: list,
    stock_chart_path: str,
    ticker: str,
    peer_chart_path: str = "",
    sentiment_summary: str = "",
    sentiment_citations: list = None,
    risk_section: str = "",
) -> str:
    """Generate a Word (.docx) report and return the file path."""

    company_name = profile.get("companyName", ticker)
    doc = Document()

    # Narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── TITLE ──
    title = doc.add_heading(company_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(20, 20, 20)

    # Subtitle
    subtitle_parts = []
    for key in ("symbol", "exchange", "sector", "industry"):
        if profile.get(key):
            subtitle_parts.append(profile[key])
    if subtitle_parts:
        sub = doc.add_paragraph(" | ".join(subtitle_parts))
        sub.style = doc.styles["No Spacing"]
        for run in sub.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(80, 80, 80)

    # Date
    date_para = doc.add_paragraph(
        f"Report generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    date_para.style = doc.styles["No Spacing"]
    for run in date_para.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Key stats
    employees = profile.get("fullTimeEmployees", "N/A")
    if employees and employees != "N/A":
        try:
            employees = f"{int(employees):,}"
        except (ValueError, TypeError):
            pass

    stats = doc.add_paragraph()
    stats.style = doc.styles["No Spacing"]
    stats_text = (
        f"Market Cap: {_fmt_number(profile.get('marketCap'), is_currency=True)}    "
        f"Price: ${profile.get('price', 'N/A')}    "
        f"Currency: {profile.get('currency', 'N/A')}    "
        f"Employees: {employees}"
    )
    run = stats.add_run(stats_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph()  # spacer

    # ── SECTION 1: COMPANY OVERVIEW ──
    _add_section_heading(doc, "1. Company Overview")
    para = doc.add_paragraph(description)
    for run in para.runs:
        run.font.size = Pt(10)

    # ── SECTION 2: BUSINESS MODEL & REVENUE BREAKDOWN ──
    _add_section_heading(doc, "2. Business Model & Revenue Breakdown")
    business_model_section = _normalize_markdown(business_model_section)

    for paragraph in business_model_section.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        # Skip if the AI repeated the section title
        stripped = paragraph.lstrip("#* ").strip()
        if stripped.lower().startswith("business model") and "revenue" in stripped.lower():
            continue
        p = doc.add_paragraph()
        # Render **bold** markdown inline
        parts = paragraph.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.size = Pt(10)
            if i % 2 == 1:
                run.bold = True
        p.paragraph_format.space_after = Pt(10)

    # ── SECTION 3: FINANCIAL SUMMARY TABLE ──
    _add_section_heading(doc, "3. Financial Summary")

    if income_statements:
        years = []
        revenues = []
        gross_profits = []
        operating_incomes = []
        net_incomes = []
        eps_values = []
        fcf_values = []

        for stmt in reversed(income_statements):
            year = stmt.get("calendarYear", stmt.get("date", "N/A")[:4])
            years.append(str(year))
            revenues.append(stmt.get("revenue"))
            gross_profits.append(stmt.get("grossProfit"))
            operating_incomes.append(stmt.get("operatingIncome"))
            net_incomes.append(stmt.get("netIncome"))
            eps_values.append(stmt.get("eps"))

        for cf in reversed(cash_flows):
            fcf = None
            ocf = cf.get("operatingCashFlow", 0) or 0
            capex = cf.get("capitalExpenditure", 0) or 0
            if ocf:
                fcf = ocf + capex
            fcf_values.append(fcf)

        while len(fcf_values) < len(years):
            fcf_values.append(None)

        rows_data = [
            ("Revenue", revenues, True),
            ("Gross Profit", gross_profits, True),
            ("Operating Income", operating_incomes, True),
            ("Net Income", net_incomes, True),
            ("EPS", eps_values, False),
            ("Free Cash Flow", fcf_values, True),
        ]

        headers = ["Metric"] + years
        table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        _style_header_row(table.rows[0], headers)

        # Data rows
        for row_idx, (label, values, is_curr) in enumerate(rows_data):
            row = table.rows[row_idx + 1]
            row.cells[0].text = label
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
            alt_color = "F0F3FA" if row_idx % 2 == 0 else "FFFFFF"
            _set_cell_shading(row.cells[0], alt_color)

            for col_idx, v in enumerate(values):
                if is_curr:
                    display = _fmt_number(v, is_currency=True)
                else:
                    display = f"{v:.2f}" if v is not None else "N/A"
                cell = row.cells[col_idx + 1]
                cell.text = display
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(8)
                _set_cell_shading(cell, alt_color)
    else:
        doc.add_paragraph(
            "Financial statement data unavailable for this ticker."
        ).italic = True

    # ── SECTION 4: STOCK PRICE CHART ──
    _add_section_heading(doc, "4. Stock Price Performance")
    if stock_chart_path and os.path.exists(stock_chart_path):
        doc.add_picture(stock_chart_path, width=Inches(6.0))
    else:
        doc.add_paragraph("Stock price chart unavailable.").italic = True

    # ── SECTION 5: COMPETITOR COMPARISON ──
    _add_section_heading(doc, "5. Competitor Comparison")

    if peer_comparison:
        headers = ["Company", "Ticker", "Mkt Cap", "P/E", "EV/EBITDA", "ROE", "Div Yield"]
        table = doc.add_table(rows=1 + len(peer_comparison), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        _style_header_row(table.rows[0], headers)

        for row_idx, peer in enumerate(peer_comparison):
            row = table.rows[row_idx + 1]
            name = peer.get("companyName", "N/A")
            if len(name) > 22:
                name = name[:20] + ".."

            row_values = [
                name,
                peer.get("ticker", "N/A"),
                _fmt_number(peer.get("mktCap"), is_currency=True),
                f"{peer.get('peRatio', 0):.1f}x" if peer.get("peRatio") else "N/A",
                f"{peer.get('evToEbitda', 0):.1f}x" if peer.get("evToEbitda") else "N/A",
                _fmt_number(peer.get("roe"), is_pct=True) if peer.get("roe") else "N/A",
                _fmt_number(peer.get("dividendYield"), is_pct=True) if peer.get("dividendYield") else "N/A",
            ]

            alt_color = "F0F3FA" if row_idx % 2 == 0 else "FFFFFF"
            for col_idx, val in enumerate(row_values):
                cell = row.cells[col_idx]
                cell.text = str(val)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(8)
                _set_cell_shading(cell, alt_color)
    else:
        doc.add_paragraph("Peer comparison data unavailable.").italic = True

    # ── SECTION 6: COMPETITIVE STOCK PERFORMANCE ──
    _add_section_heading(doc, "6. Competitive Stock Performance")
    if peer_chart_path and os.path.exists(peer_chart_path):
        doc.add_picture(peer_chart_path, width=Inches(6.0))
    else:
        doc.add_paragraph("Competitive price chart unavailable.").italic = True

    # ── SECTION 7: KEY RISKS ──
    if risk_section:
        _add_section_heading(doc, "7. Key Risks")
        risk_section = _normalize_markdown(risk_section)
        for paragraph in risk_section.split("\n\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            stripped = paragraph.lstrip("#* ").strip()
            if stripped.lower().startswith("key risk"):
                continue
            p = doc.add_paragraph()
            parts = paragraph.split("**")
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                if i % 2 == 1:
                    run.bold = True
            p.paragraph_format.space_after = Pt(10)

    # ── SECTION 8: MARKET SENTIMENT ──
    if sentiment_summary:
        _add_section_heading(doc, "8. Market Sentiment and Recent News")
        para = doc.add_paragraph(sentiment_summary)
        for run in para.runs:
            run.font.size = Pt(10)

        if sentiment_citations:
            sources_para = doc.add_paragraph()
            run = sources_para.add_run("Sources:")
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(80, 80, 80)

            for i, url in enumerate(sentiment_citations, 1):
                display = url if len(url) <= 90 else url[:87] + "..."
                p = doc.add_paragraph()
                p.style = doc.styles["No Spacing"]
                r = p.add_run(f"[{i}]  {display}")
                r.font.size = Pt(7)
                r.font.italic = True
                r.font.color.rgb = RGBColor(100, 100, 100)

    # ── DISCLAIMER FOOTER ──
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "This report is auto-generated for informational purposes only. "
        "It does not constitute investment advice."
    )
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)

    # ── SAVE ──
    output_path = os.path.join(tempfile.gettempdir(), f"{ticker}_report.docx")
    doc.save(output_path)
    return output_path
